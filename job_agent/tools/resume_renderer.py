"""
Renders a tailored resume from a DOCX template.

Template lives at: data/profile/resume_template.docx
Generated resumes go to: data/resumes/{company}_{title}.docx

ATS note: the template must use plain paragraphs only.
No tables, no text boxes, no columns — ATS parsers reject them.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document

TEMPLATE_PATH = Path("data/profile/resume_template.docx")
OUTPUT_DIR    = Path("data/resumes")


def render(replacements: dict[str, str], output_filename: str) -> Path:
    """
    Fill all {{key}} placeholders in the DOCX template and save.

    Parameters
    ----------
    replacements    : mapping of placeholder → value
    output_filename : e.g. "stripe_senior_engineer.docx"

    Returns the path to the saved file.
    """
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Resume template not found at {TEMPLATE_PATH}")

    doc = Document(TEMPLATE_PATH)

    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = OUTPUT_DIR / output_filename
    doc.save(out_path)
    return out_path


def _replace_in_paragraph(para, replacements: dict[str, str]) -> None:
    """Replace {{key}} placeholders while preserving run-level formatting."""
    for key, value in replacements.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in para.text:
            for run in para.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)
